import os
import json
import sqlite3
from pathlib import Path
from typing import List, Tuple


def extract_pdf(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber required: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"[PAGE {page_num}]\n{text}")
    return "\n\n".join(text_parts)


def extract_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text_from_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[dict]:
    chunks = []
    sentences = text.split(". ")
    current_chunk = ""
    chunk_index = 0

    for sentence in sentences:
        if len(current_chunk) + len(sentence) > chunk_size:
            if current_chunk.strip():
                chunks.append({
                    "chunk_index": chunk_index,
                    "content": current_chunk.strip(),
                    "location_hint": _extract_location_hint(current_chunk),
                })
                chunk_index += 1
                current_chunk = current_chunk[-overlap:] + ". " if len(current_chunk) > overlap else ""
        current_chunk += sentence + ". "

    if current_chunk.strip():
        chunks.append({
            "chunk_index": chunk_index,
            "content": current_chunk.strip(),
            "location_hint": _extract_location_hint(current_chunk),
        })
    return chunks


def _extract_location_hint(text: str) -> str:
    if "[PAGE" in text:
        for part in text.split("[PAGE"):
            if "]" in part:
                return f"Page {part.split(']')[0].strip()}"
    return "Unknown"


def ingest_case_file(case_id: int, file_path: str) -> Tuple[int, List[dict]]:
    """
    Extract, chunk, embed, and store all chunks.
    Returns (num_chunks, chunk_list).
    """
    text = extract_text_from_file(file_path)
    if not text or len(text.strip()) < 100:
        raise ValueError("Extracted text is too short or empty.")

    chunks = chunk_text(text, chunk_size=1000, overlap=200)

    # Embed all chunks at once
    from llm_client import get_llm_client
    llm = get_llm_client()
    contents = [c["content"] for c in chunks]
    embeddings = llm.embed(contents)

    db_path = "data/db/casesim.db"
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Clear old chunks for this case before re-ingesting
    cursor.execute("DELETE FROM case_chunks WHERE case_id = ?", (case_id,))

    for chunk, embedding in zip(chunks, embeddings):
        cursor.execute("""
            INSERT INTO case_chunks (case_id, chunk_index, location_hint, content, embedding_json)
            VALUES (?, ?, ?, ?, ?)
        """, (
            case_id,
            chunk["chunk_index"],
            chunk["location_hint"],
            chunk["content"],
            json.dumps(embedding),
        ))

    conn.commit()
    conn.close()
    return len(chunks), chunks


def get_case_chunks(case_id: int, limit: int = 200) -> List[dict]:
    db_path = "data/db/casesim.db"
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, chunk_index, location_hint, content
        FROM case_chunks WHERE case_id = ?
        ORDER BY chunk_index LIMIT ?
    """, (case_id, limit))
    chunks = [
        {"chunk_id": r[0], "chunk_index": r[1], "location_hint": r[2], "content": r[3]}
        for r in cursor.fetchall()
    ]
    conn.close()
    return chunks
